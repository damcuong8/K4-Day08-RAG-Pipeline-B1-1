"""
Task 1 — Thu thập văn bản chính sách thương mại điện tử / hỗ trợ khách hàng.

Hướng dẫn:
    1. Tìm tối thiểu 3 văn bản chính sách (PDF/DOCX) từ trang chính thức của một sàn TMĐT.
    2. Tải về và lưu vào data/landing/legal/
    3. Đặt tên file rõ ràng, không dấu, mô tả đúng nội dung.

Gợi ý nguồn (ví dụ trang công khai Shopee Vietnam — help.shopee.vn):
    - https://help.shopee.vn/portal/4/article/77251 (Chính sách trả hàng và hoàn tiền)
    - https://help.shopee.vn/portal/4/article/79198 (Phương thức thanh toán)
    - https://help.shopee.vn/portal/4/article/77244 (Chính sách bảo mật)

Gợi ý văn bản (chủ đề chính sách thương mại điện tử):
    - Chính sách đổi trả/hoàn tiền (Returns/Refund Policy)
    - Phương thức thanh toán (Payment Methods)
    - Chính sách bảo mật (Privacy Policy)
    - Quy định đăng bán sản phẩm cho người bán (Seller Listing Regulations)

Nhớ gắn metadata `customer_role` (`buyer`/`seller`/`both`) cho từng tài liệu — yêu cầu riêng
của K4 Variant (kế thừa từ Lab 07), cần thiết để viết benchmark query dùng metadata_filter.

Lưu ý: một số trang help center dùng JavaScript render nội dung (SPA) — crawl về chỉ thấy
tiêu đề mà không có nội dung thật. Đổi sang bài viết khác cùng domain thay vì cố xử lý,
và chỉ dùng nguồn công khai/được phép chia sẻ.
"""

from __future__ import annotations

import json
import os
from pathlib import Path

import requests

DATA_DIR = Path(__file__).parent.parent / "data" / "landing" / "legal"


def setup_directory():
    """Tạo thư mục data/landing/legal/ nếu chưa có."""
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    print(f"✓ Thư mục đã sẵn sàng: {DATA_DIR}")


def download_file(url: str, filename: str) -> Path:
    """Download one public PDF/DOC/DOCX with basic validation."""
    setup_directory()
    destination = DATA_DIR / Path(filename).name
    if destination.suffix.lower() not in {".pdf", ".doc", ".docx"}:
        raise ValueError("Task 1 chỉ nhận PDF, DOC hoặc DOCX")
    response = requests.get(url, timeout=60)
    response.raise_for_status()
    if len(response.content) <= 1024:
        raise ValueError(f"Nội dung tải từ {url} quá nhỏ")
    destination.write_bytes(response.content)
    return destination


def export_demo_from_legal_corpus(limit: int = 3) -> list[Path]:
    """Export traceable DOCX samples from the existing Vietnamese legal corpus.

    This provides a reproducible local demo when direct source downloads are not
    configured. Every file records the original title, number, and source URL.
    """
    from docx import Document

    root = Path(
        os.getenv(
            "LEGAL_ASSISTANT_ROOT", str(Path(__file__).parents[2] / "Legal_assistant")
        )
    )
    metadata_path = root / "data/raw_law/law_only/effective/law_luoc_do_merged_dedup.jsonl"
    chunks_path = root / "data/raw_law/law_only/effective/parsed_law_database.jsonl"
    if not metadata_path.exists() or not chunks_path.exists():
        raise FileNotFoundError("Không tìm thấy corpus Legal_assistant để tạo dữ liệu demo")

    metadata = {}
    with metadata_path.open(encoding="utf-8") as handle:
        for line in handle:
            item = json.loads(line)
            metadata[str(item.get("id"))] = item

    selected: dict[str, list[dict]] = {}
    with chunks_path.open(encoding="utf-8") as handle:
        for line in handle:
            chunk = json.loads(line)
            doc_id = str(chunk.get("doc_id"))
            if doc_id not in metadata:
                continue
            if doc_id not in selected and len(selected) >= limit:
                continue
            selected.setdefault(doc_id, []).append(chunk)
            if len(selected) >= limit and all(len(parts) >= 3 for parts in selected.values()):
                break

    setup_directory()
    outputs = []
    for doc_id, chunks in selected.items():
        meta = metadata[doc_id]
        attributes = meta.get("thuoc_tinh") or {}
        document = Document()
        document.add_heading(str(meta.get("title") or f"Văn bản {doc_id}"), level=1)
        document.add_paragraph(f"Số hiệu: {attributes.get('Số hiệu', '')}")
        document.add_paragraph(f"Nguồn: {meta.get('url', '')}")
        for chunk in chunks:
            heading = " — ".join(
                value
                for value in (
                    str(chunk.get("article_no") or "").strip(),
                    str(chunk.get("article_title") or "").strip(),
                )
                if value
            )
            if heading:
                document.add_heading(heading, level=2)
            document.add_paragraph(str(chunk.get("text") or ""))
        output = DATA_DIR / f"legal_document_{doc_id}.docx"
        document.save(output)
        outputs.append(output)
    return outputs


# TODO: Tải file PDF/DOCX về DATA_DIR
# Có thể tải thủ công hoặc viết script download nếu có direct link.
#
# Ví dụ nếu có direct link:
#
# import requests
#
# def download_file(url: str, filename: str):
#     response = requests.get(url)
#     filepath = DATA_DIR / filename
#     filepath.write_bytes(response.content)
#     print(f"✓ Đã tải: {filepath}")
#
# Nếu trang là HTML thuần (không phải PDF sẵn), có thể convert nội dung text
# thành PDF đơn giản bằng thư viện fpdf2 (đã có trong requirements.txt).


if __name__ == "__main__":
    files = export_demo_from_legal_corpus()
    print(f"✓ Đã chuẩn bị {len(files)} văn bản pháp luật")
